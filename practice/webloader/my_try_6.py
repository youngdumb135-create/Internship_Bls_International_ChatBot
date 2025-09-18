from pdf2docx import Converter
from docx import Document
import os

pdf_file = r"D:\Chatbot\practice\pdfs\APPLICATION_FOR_SCHENGEN_VISA_English.pdf"
docx_file = "D:\Chatbot\practice\pdfs\APPLICATION_FOR_SCHENGEN_VISA_English.docx"


def pdf_to_docx_conversion(pdf_location, docx_location):
    if not os.path.exists(pdf_location):
        print(f"Error: PDF file not found at {pdf_location}")
        return None
    
    cv = Converter(pdf_file)
    cv.convert(docx_file, start=0, end=None)  
    cv.close()
    print(f"Successfully converted {pdf_file} to {docx_file} using Converter class.")

    return docx_location


def extract_content_from_docx(docx_file):
    if not os.path.exists(docx_file):
        print(f"Error: DOCX file not found at {docx_file}")
        return []

    document = Document(docx_file)
    extracted_data = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            extracted_data.append({"type": "text", "content": paragraph.text})

    for table in document.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)

        if table_data:
            extracted_data.append({"type": "table", "content": table_data})

    return extracted_data




if __name__ == "__main__":
    pdf_file = r"D:\Chatbot\practice\pdfs\APPLICATION_FOR_SCHENGEN_VISA_English.pdf"
    docx_file = r"D:\Chatbot\practice\pdfs\APPLICATION_FOR_SCHENGEN_VISA_English.docx"

    docx_path = pdf_to_docx_conversion(pdf_file, docx_file)
    a = extract_content_from_docx(docx_path)
    print(a)


































# import pdfplumber
# import pandas as pd

# def extract_separated_content_as_markdown(pdf_path):
#     """
#     Extracts and separates text and tables from a PDF, formatting the output
#     with Markdown syntax.

#     Args:
#         pdf_path (str): The path to the PDF file.
    
#     Returns:
#         str: A Markdown-formatted string of the extracted content.
#     """
#     markdown_output = []
    
#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             markdown_output.append(f"# Content from: {pdf_path}\n")

#             for i, page in enumerate(pdf.pages):
#                 markdown_output.append(f"## Page {i+1}\n")

#                 # Step 1: Find tables and their bounding boxes
#                 tables = page.extract_tables()
#                 table_bboxes = [table_obj.bbox for table_obj in page.find_tables()]

#                 # Step 2: Extract and format tables in Markdown
#                 if tables:
#                     for table_num, table in enumerate(tables):
#                         markdown_output.append(f"### Table {table_num + 1}\n")
#                         df = pd.DataFrame(table[1:], columns=table)
#                         markdown_output.append(df.to_markdown(index=False, tablefmt="github"))
#                         markdown_output.append("\n")
                
#                 # Step 3: Create a version of the page that excludes table areas
#                 # This ensures the text within tables is not extracted twice.
#                 cropped_page = page
#                 for bbox in table_bboxes:
#                     cropped_page = cropped_page.outside_bbox(bbox)
                
#                 # Step 4: Extract text from the non-table areas
#                 text = cropped_page.extract_text()
                
#                 markdown_output.append("### Text Content (Outside Tables)\n")
#                 if text:
#                     markdown_output.append(text)
#                 else:
#                     markdown_output.append("No text found outside of tables on this page.")
                
#                 markdown_output.append("\n---\n")

#     except FileNotFoundError:
#         return f"Error: The file at '{pdf_path}' was not found."
#     except Exception as e:
#         return f"An error occurred: {e}"

#     return "\n".join(markdown_output)

# # Example usage:
# file_name = 'D:\Chatbot\practice\pdfs\APPLICATION_FOR_SCHENGEN_VISA_English.pdf'
# extracted_content = extract_separated_content_as_markdown(file_name)
# print(extracted_content)

